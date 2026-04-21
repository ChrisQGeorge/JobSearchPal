"""Best-effort text extraction for common document types.

Called from the upload endpoint so PDFs / DOCX / HTML files get a populated
`content_md` alongside their preserved binary. That lets both the in-app
editor and the Companion work on the text content directly.

Failures are intentionally non-fatal — if extraction returns None the upload
still succeeds, the user just can't edit the text content in-app.
"""

from __future__ import annotations

import io
import logging
from typing import Optional

log = logging.getLogger(__name__)


# Text cap so pathological uploads don't bloat the DB. 2 MB of extracted text
# is ~2M chars / ~400k words — far more than any sane resume, cover letter, or
# single-doc upload. Above this we truncate with a marker.
MAX_EXTRACTED_CHARS = 2_000_000


def _is_pdf(mime: str, filename: str) -> bool:
    return mime == "application/pdf" or filename.lower().endswith(".pdf")


def _is_docx(mime: str, filename: str) -> bool:
    # .docx is the Office Open XML format; older .doc binary format is
    # intentionally unsupported (would need antiword or similar).
    return (
        mime
        in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",  # some browsers send this for .docx too
        )
        or filename.lower().endswith(".docx")
    )


def _is_html(mime: str, filename: str) -> bool:
    if mime.startswith("text/html") or mime == "application/xhtml+xml":
        return True
    return filename.lower().endswith((".html", ".htm", ".xhtml"))


def _clamp(text: str) -> str:
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text
    return text[:MAX_EXTRACTED_CHARS] + "\n\n[… truncated, extracted text exceeded 2 MB …]"


def _extract_pdf(data: bytes) -> Optional[str]:
    try:
        from pypdf import PdfReader
    except ImportError:  # pragma: no cover
        log.warning("pypdf not available; skipping PDF extraction")
        return None
    try:
        reader = PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                pages.append(page.extract_text() or "")
            except Exception as exc:  # pragma: no cover
                log.warning("PDF page %d extraction failed: %s", i, exc)
                pages.append("")
        text = "\n\n".join(p.strip() for p in pages if p and p.strip())
        return _clamp(text) if text else None
    except Exception as exc:
        log.warning("PDF extraction failed: %s", exc)
        return None


def _extract_docx(data: bytes) -> Optional[str]:
    try:
        from docx import Document
    except ImportError:  # pragma: no cover
        log.warning("python-docx not available; skipping DOCX extraction")
        return None
    try:
        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
        # Preserve table content — resumes and offer letters often use them.
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts)
        return _clamp(text) if text else None
    except Exception as exc:
        log.warning("DOCX extraction failed: %s", exc)
        return None


def _extract_html(data: bytes) -> Optional[str]:
    # Decode: prefer UTF-8, fall back to latin-1 (never errors).
    try:
        raw = data.decode("utf-8")
    except UnicodeDecodeError:
        raw = data.decode("latin-1", errors="replace")

    # Strip script/style blocks via BeautifulSoup first, then convert to
    # Markdown so the result is editor-friendly.
    try:
        from bs4 import BeautifulSoup
        import html2text
    except ImportError:  # pragma: no cover
        log.warning("bs4/html2text not available; skipping HTML extraction")
        return None
    try:
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        cleaned = str(soup)
        h = html2text.HTML2Text()
        h.body_width = 0  # don't re-wrap lines
        h.ignore_images = False
        h.ignore_links = False
        text = h.handle(cleaned).strip()
        return _clamp(text) if text else None
    except Exception as exc:
        log.warning("HTML extraction failed: %s", exc)
        return None


def extract_text(data: bytes, mime: str, filename: str) -> Optional[str]:
    """Return extracted markdown/plain-text for supported types, else None.

    Recognizes: plain text (text/*, .md, .txt), PDF, DOCX, HTML. Returns
    None for types we don't know how to read.
    """
    mime = (mime or "").lower()
    name_lower = filename.lower()

    # Plain text — decode directly. `text/*`, JSON, .md, .txt.
    is_plain = (
        mime.startswith("text/")
        and not mime.startswith("text/html")
        or mime == "application/json"
        or name_lower.endswith((".md", ".txt", ".markdown"))
    )
    if is_plain:
        try:
            return _clamp(data.decode("utf-8"))
        except UnicodeDecodeError:
            return None

    if _is_pdf(mime, filename):
        return _extract_pdf(data)
    if _is_docx(mime, filename):
        return _extract_docx(data)
    if _is_html(mime, filename):
        return _extract_html(data)

    return None


def kind_of(mime: str, filename: str) -> str:
    """Short label for the kind of source a given upload is. Used in
    `content_structured.extracted_from` so the UI can show 'Extracted from PDF'
    and so the Companion knows the original format."""
    if _is_pdf(mime, filename):
        return "pdf"
    if _is_docx(mime, filename):
        return "docx"
    if _is_html(mime, filename):
        return "html"
    if mime.startswith("text/") or filename.lower().endswith((".md", ".txt", ".markdown")):
        return "text"
    return "binary"
